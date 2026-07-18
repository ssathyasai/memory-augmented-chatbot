"""File parsers for different document formats."""

import logging
from typing import Optional
from io import BytesIO
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument

from errors.exceptions import DocumentProcessingError

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse documents of various formats."""
    
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """
        Parse PDF file.
        
        Args:
            file_bytes: PDF file bytes
        
        Returns:
            Extracted text content
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                if text.strip():
                    return text.strip()
            
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
            
            if not text.strip():
                raise DocumentProcessingError("Could not extract text from PDF")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise DocumentProcessingError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """
        Parse DOCX file.
        
        Args:
            file_bytes: DOCX file bytes
        
        Returns:
            Extracted text content
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        try:
            doc = DocxDocument(BytesIO(file_bytes))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise DocumentProcessingError("DOCX file is empty")
            
            return text.strip()
        
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """
        Parse TXT file.
        
        Args:
            file_bytes: TXT file bytes
        
        Returns:
            Text content
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_bytes.decode(encoding)
                    if text.strip():
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError("Could not decode text file")
        
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise DocumentProcessingError(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    def parse_markdown(file_bytes: bytes) -> str:
        """
        Parse Markdown file.
        
        Args:
            file_bytes: Markdown file bytes
        
        Returns:
            Text content
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        # Markdown is essentially text, so use same logic
        return DocumentParser.parse_txt(file_bytes)
    
    @staticmethod
    def parse(file_bytes: bytes, file_type: str) -> str:
        """
        Parse file based on type.
        
        Args:
            file_bytes: File bytes
            file_type: File extension (.pdf, .docx, .txt, .md)
        
        Returns:
            Extracted text content
        
        Raises:
            DocumentProcessingError: If parsing fails
        """
        file_type = file_type.lower().replace('.', '')
        
        if file_type == 'pdf':
            return DocumentParser.parse_pdf(file_bytes)
        elif file_type == 'docx':
            return DocumentParser.parse_docx(file_bytes)
        elif file_type == 'txt':
            return DocumentParser.parse_txt(file_bytes)
        elif file_type in ['md', 'markdown']:
            return DocumentParser.parse_markdown(file_bytes)
        else:
            raise DocumentProcessingError(f"Unsupported file type: {file_type}")
